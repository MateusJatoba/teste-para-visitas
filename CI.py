from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

class ComunicacaoInterna:
    def __init__(self):
        self.doc = Document()
        self.save_path = r'C:\Users\mateus.freire\OneDrive - Sistema FIEB\Documentos\prospecção\BD_agenda'
        self.borda_padrao = {
            'val': 'single',
            'sz': '6',
            'space': '0',
            'color': '000000'
        }

    def set_cell_text(self, cell, text, bold=False):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        run.font.name = "Arial"

    def set_cell_border(self, cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for edge in ('top', 'bottom', 'start', 'end'):
            edge_data = kwargs.get(edge)
            if edge_data:
                element = OxmlElement(f'w:{edge}')
                element.set(qn('w:val'), edge_data.get('val', 'single'))
                element.set(qn('w:sz'), edge_data.get('sz', '6'))
                element.set(qn('w:space'), edge_data.get('space', '0'))
                element.set(qn('w:color'), edge_data.get('color', '000000'))
                tcPr.append(element)

    def gerar_documento(self):
        doc = self.doc
        doc.add_paragraph("COMUNICAÇÃO INTERNA").runs[0].bold = True

        table = doc.add_table(rows=3, cols=4)
        table.autofit = False

        for i, width in enumerate([Pt(80), Pt(200), Pt(80), Pt(200)]):
            table.columns[i].width = width

        # Linha 1
        self.set_cell_text(table.cell(0, 0), "DATA", bold=True)
        self.set_cell_text(table.cell(0, 1), "15/05/2025")
        self.set_cell_text(table.cell(0, 2), "Nº:", bold=True)
        self.set_cell_text(table.cell(0, 3), "xxx/2024")

        # Linha 2
        self.set_cell_text(table.cell(1, 0), "DE:", bold=True)
        self.set_cell_text(table.cell(1, 1), "Marina Peixoto Santos\nGerente de Negócios SENAI CIMATEC")
        self.set_cell_text(table.cell(1, 2), "PARA:", bold=True)
        self.set_cell_text(table.cell(1, 3), "Leone Andrade\nDiretor Geral SENAI CIMATEC")

        # Linha 3
        cell_assunto = table.cell(2, 0).merge(table.cell(2, 1)).merge(table.cell(2, 2)).merge(table.cell(2, 3))
        self.set_cell_text(cell_assunto, "ASSUNTO:  Autorização para Contratação de Alimentos e Bebidas", bold=True)

        # Bordas
        for row in table.rows:
            for cell in row.cells:
                self.set_cell_border(cell, top=self.borda_padrao, bottom=self.borda_padrao,
                                     start=self.borda_padrao, end=self.borda_padrao)

        # Texto principal
        doc.add_paragraph("\nPrezado Diretor,")
        doc.add_paragraph(
            "Solicitamos autorização para contratação de pratos gourmet no restaurante Páprica no CIMATEC PARK para fornecimento de almoço de representação, incluindo mão de obra (caso necessário)."
        )
        doc.add_paragraph("Seguem informações detalhadas:")

        info = {
            "Empresa": f"{self.extrai_visita()[0]}",
            "Justificativa com o objetivo da visita": (
                "A visita da equipe da Petrobras às nossas instalações tem como principal finalidade acompanhar de perto "
                "a demonstração do projeto SDM, permitindo uma visão mais detalhada das soluções desenvolvidas e das tecnologias aplicadas. "
                "Esta será uma oportunidade para apresentar os avanços alcançados, discutir aspectos técnicos e operacionais, "
                "além de alinhar expectativas quanto aos próximos passos do projeto. "
                "A visita contempla 2 dias de demonstrações e esse pedido contempla o segundo dia que não estava contemplado na CI 1294."
            ),
            "Data": "15/05/2025",
            "Horário": "8h às 14h",
            "Quantidade de participantes externos": "4",
            "Quantidade de participantes internos": "0",
            "Valor estimando por pessoa": "R$ 80,00",
            "Valor estimando total": "R$ 320,00",
            "CR para alocar as despesas": "30202010010077140010 PJ PETROBRAS SUB MERG ANP-IPD(ROB)CONTR 30048 ROBÓTICA"
        }

        for key, value in info.items():
            doc.add_paragraph(f"{key}: {value}")

        doc.add_paragraph("\nPor meio deste, requeremos sua aprovação para que possamos prosseguir com a solicitação.")
        doc.add_paragraph("Atenciosamente,")
        doc.add_paragraph("Marina Peixoto Santos")
        doc.add_paragraph("Gerente de Negócios\nRobótica\nSENAI CIMATEC")

        # Salvar
        filename = os.path.join(self.save_path, "CI_xxxx_2024_Solicitacao_Alimentacao.docx")
        doc.save(filename)
        print(f"Documento salvo em: {filename}")


# Rodar
ci = ComunicacaoInterna()
ci.gerar_documento()

