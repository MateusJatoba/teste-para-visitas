from docx import Document  
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import conexao_ollama
import sqlite3
from time import sleep

class Agenda:
    def __init__(self):    
        self.doc = self.criar_documento()
        self.save_path = r'C:\Users\mateus.freire\OneDrive - Sistema FIEB\Documentos\prospecção\BD_agenda' 
    
    
    def criar_documento(self):
        return Document()
    
    def extract_data(self):
        conn = sqlite3.connect("banco_visitas.db")
        cursor = conn.cursor()

        cursor.execute("SELECT MAX(id) FROM visita_participante_Interno")
        last_id_created = cursor.fetchone()[0]

        cursor.execute(f"SELECT (id_visita) FROM visita_participante_Interno WHERE id = {last_id_created}")
        last_visit = cursor.fetchone()[0]

        cursor.execute(f"SELECT * FROM TB_visita WHERE id = {last_visit}")
        last_visit_data = cursor.fetchall()

        cursor.execute("""
            SELECT p.* 
            FROM TB_participante_Interno p
            JOIN visita_participante_Interno vpi ON p.id = vpi.id_participante
            WHERE vpi.id_visita = ?
        """, (last_visit,)
        )

        people = cursor.fetchall()

        cursor.execute("""
            SELECT * FROM TB_participante_Externo
            WHERE id_visita = (
                SELECT MAX(id_visita) FROM TB_participante_Externo
            );
        """)
        visitors = cursor.fetchall()

        
        conn.close()


        return [people , last_visit_data, visitors]
            
      

    def add_hyperlink(self, paragraph, url, text, color="0000FF", underline=True):
        """Função para inserir hiperlink no Word usando python-docx."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

      
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), color)
        rPr.append(color_elem)

   
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink


    def manipula_documento(self):
        
        paragrafo1 = self.doc.add_paragraph()
        imagem = paragrafo1.add_run().add_picture(r'C:\Users\mateus.freire\OneDrive - Sistema FIEB\Documentos\prospecção\pasta_auxiliar\logo_CIMATEC.png', width=Inches(2))

        paragrafo1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = self.doc.add_paragraph().add_run(
            "O SENAI CIMATEC é um dos mais avançados centros de tecnologia e inovação do Brasil, especializado em desenvolver "
            "pesquisas e soluções para a indústria e para Micro, Pequenas e Médias Empresas Industriais. "
            "Com sede em Salvador e um time de mais de 1.000 pessoas, a instituição, sem fins lucrativos, integra Centros Tecnológico, "
            "Educação Profissional e um Centro Universitário, com cursos de graduação, especializações, Mestrado e Doutorado. "
            "O Centro Tecnológico possui 44 áreas de competência, entre elas: Robótica e Automação, Energia e Sustentabilidade, "
            "Saúde, Alimentos, Software e Supercomputação. "
            "Em 2019, foi inaugurado o SENAI CIMATEC Park, complexo de inovação industrial que reúne os conceitos de parques "
            "científico, tecnológico e de negócios, em uma área de 4 milhões de m², no Polo Industrial de Camaçari."
        )
        run.font.name = 'Aptos'
        run.font.size = Pt(9)


        p2 = self.doc.add_paragraph()
        titulo = p2.add_run(f'{self.extract_data()[1][0][1].upper()}')
        titulo.font.name = 'Aptos Display Bold'
        titulo.font.size = Pt(12)
        titulo.font.bold = True
        
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        
        # p2b = self.doc.add_paragraph()
        # run2b = p2b.add_run(
        #     conexao_ollama.conversar(self.extract_data()[1][0][1])
        # )
        # run2b.font.name = 'Aptos'
        # run2b.font.size = Pt(12)

        # Participantes
        p_participantes = self.doc.add_paragraph()
        p_participantes.add_run("PARTICIPANTES INTERNOS:\n").bold = True
        p_participantes.add_run(
            ' / '.join(f'{i[2]} ({i[3]})' for i in a.extract_data()[0])
        )
        p_participantes.paragraph_format.space_after = Pt(6)

        p_participantes_externos = self.doc.add_paragraph()
        p_participantes_externos.add_run("PARTICIPANTES EXTERNOS:\n").bold = True
        p_participantes_externos.add_run(
            '\n'.join(f'{i[2]} - {i[3]}' for i in a.extract_data()[2])
        )




        if (self.extract_data()[1][0][3] == '1'):

            p3 = self.doc.add_paragraph()
            p3.add_run("\nSENAI CIMATEC\n").bold = True
            p3.add_run("Av. Orlando Gomes, 1845 – Piatã, Salvador\nPrédio I, 2º andar, Auditório Gerência\nLocalização: ")

            self.add_hyperlink(p3, "https://goo.gl/maps/yGD3UbucB4jCt9v27", "https://goo.gl/maps/yGD3UbucB4jCt9v27")
        
        if (self.extract_data()[1][0][3] == '2'):
            p3 = self.doc.add_paragraph()
            p3.add_run("\n\nCIMATEC PARK\n").bold = True
            p3.add_run("Via Atlântica, BA 530, KM 2,5 - Camaçari\nLocalização: ")

            self.add_hyperlink(p3, "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8", "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8")
        
        if (self.extract_data()[1][0][3] == '3'):
            p3 = self.doc.add_paragraph()
            p3.add_run("\nSENAI CIMATEC\n").bold = True
            p3.add_run("Av. Orlando Gomes, 1845 - Piatã, Salvador\nPrédio I, 2º andar, Auditório Gerência\nLocalização: ")

            self.add_hyperlink(p3, "https://goo.gl/maps/yGD3UbucB4jCt9v27", "https://goo.gl/maps/yGD3UbucB4jCt9v27")

            p3.add_run("\n\nCIMATEC PARK\n").bold = True
            p3.add_run("Via Atlântica, BA 530, KM 2,5 - Camaçari\nLocalização: ")

            self.add_hyperlink(p3, "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8", "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8")
        
        p_agenda_titulo = self.doc.add_paragraph()

        texto_roteiro = self.extract_data()[1][0][7]
        texto_roteiro = texto_roteiro.replace('\r\n', '\n').replace('\r', '\n')


        p_agenda_titulo.add_run(f"{texto_roteiro}") # roteiro da visita 

        
        file_path = os.path.join(self.save_path, f'AGENDA_{self.extract_data()[1][0][1]}_2025.docx')
        self.doc.save(file_path)
        
sleep(4)
a = Agenda()
# for visitors in a.extract_data()[2]:
#     print(visitors[2])
a.manipula_documento()

# print(a.extract_data()[2])

# print(a.extract_data()[1][0][1]) 

# print(' / '.join(f'{i[2]} ({i[3]})' for i in a.extract_data()[0]))


# print(f'{a.extract_data()[0]} ')

