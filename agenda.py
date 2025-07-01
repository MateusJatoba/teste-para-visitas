from docx import Document  
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import openpyxl
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import conexao_ollama


class Agenda:
    def __init__(self):    
        self.doc = self.criar_documento()
        self.save_path = r'C:\Users\mateus.freire\OneDrive - Sistema FIEB\Documentos\prospecção\BD_agenda' 
    
    
    def criar_documento(self):
        return Document()
    
    def extrai_visita(self):
        df = pd.read_excel(r'C:\Users\mateus.freire\Sistema FIEB\Testes Para Novas Ações - Documentos\teste_prospeccao\LOG_PROSPEC.xlsx', header=0, engine="openpyxl")
        # print(df) 
        return [df["Título"].iloc[-1], df["Endereço da Visita"].iloc[-1], df["Descrição da Agenda"].iloc[-1], df["Carro/van"].iloc[-1]]
      

    def add_hyperlink(self, paragraph, url, text, color="0000FF", underline=True):
        """Função para inserir hiperlink no Word usando python-docx."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

      
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), color)
        rPr.append(color_elem)

   
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink


    def manipula_documento(self):
        
        paragrafo1 = self.doc.add_paragraph()
        imagem = paragrafo1.add_run().add_picture(r'C:\Users\mateus.freire\OneDrive - Sistema FIEB\Documentos\prospecção\pasta_auxiliar\logo_CIMATEC.png', width=Inches(2))

        paragrafo1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = self.doc.add_paragraph().add_run(
            "O SENAI CIMATEC é um dos mais avançados centros de tecnologia e inovação do Brasil, especializado em desenvolver "
            "pesquisas e soluções para a indústria e para Micro, Pequenas e Médias Empresas Industriais. "
            "Com sede em Salvador e um time de mais de 1.000 pessoas, a instituição, sem fins lucrativos, integra Centros Tecnológico, "
            "Educação Profissional e um Centro Universitário, com cursos de graduação, especializações, Mestrado e Doutorado. "
            "O Centro Tecnológico possui 44 áreas de competência, entre elas: Robótica e Automação, Energia e Sustentabilidade, "
            "Saúde, Alimentos, Software e Supercomputação. "
            "Em 2019, foi inaugurado o SENAI CIMATEC Park, complexo de inovação industrial que reúne os conceitos de parques "
            "científico, tecnológico e de negócios, em uma área de 4 milhões de m², no Polo Industrial de Camaçari."
        )
        run.font.name = 'Aptos'
        run.font.size = Pt(9)


        p2 = self.doc.add_paragraph()
        titulo = p2.add_run(f'{self.extrai_visita()[0]}')
        titulo.font.name = 'Aptos Display Bold'
        titulo.font.size = Pt(12)
        titulo.font.bold = True
        
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        
        p2b = self.doc.add_paragraph()
        run2b = p2b.add_run(
            conexao_ollama.conversar(self.extrai_visita()[0])
        )
        run2b.font.name = 'Aptos'
        run2b.font.size = Pt(12)



        if (self.extrai_visita()[1] == 'Cimatec Piatã'):

            p3 = self.doc.add_paragraph()
            p3.add_run("\nSENAI CIMATEC\n").bold = True
            p3.add_run("Av. Orlando Gomes, 1845 – Piatã, Salvador\nPrédio I, 2º andar, Auditório Gerência\nLocalização: ")

            self.add_hyperlink(p3, "https://goo.gl/maps/yGD3UbucB4jCt9v27", "https://goo.gl/maps/yGD3UbucB4jCt9v27")
        
        if (self.extrai_visita()[1] == 'Cimatec Park'):
            p3 = self.doc.add_paragraph()
            p3.add_run("\n\nCIMATEC PARK\n").bold = True
            p3.add_run("Via Atlântica, BA 530, KM 2,5 - Camaçari\nLocalização: ")

            self.add_hyperlink(p3, "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8", "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8")
        
        if (self.extrai_visita()[1] == 'Mista'):
            p3 = self.doc.add_paragraph()
            p3.add_run("\nSENAI CIMATEC\n").bold = True
            p3.add_run("Av. Orlando Gomes, 1845 – Piatã, Salvador\nPrédio I, 2º andar, Auditório Gerência\nLocalização: ")

            self.add_hyperlink(p3, "https://goo.gl/maps/yGD3UbucB4jCt9v27", "https://goo.gl/maps/yGD3UbucB4jCt9v27")

            p3.add_run("\n\nCIMATEC PARK\n").bold = True
            p3.add_run("Via Atlântica, BA 530, KM 2,5 - Camaçari\nLocalização: ")

            self.add_hyperlink(p3, "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8", "https://maps.app.goo.gl/bHrpV4Drbf48oZHF8")
        
        p_agenda_titulo = self.doc.add_paragraph()

        p_agenda_titulo.add_run(f"{self.extrai_visita()[2]}")

        
        file_path = os.path.join(self.save_path, f'AGENDA_{self.extrai_visita()[0]}_2025.docx')
        self.doc.save(file_path)
        

a = Agenda()
# a.manipula_documento()
# print(a.extrai_visita())

